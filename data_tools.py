"""데이터 분석 에이전트용 Tool 모듈.

AgentBot과 함께 사용할 수 있는 데이터 분석 도구들을 제공합니다.
CSV 파일을 로드하고, Pandas 쿼리를 실행하고, 컬럼 통계를 조회하고,
차트를 생성할 수 있습니다.
"""

import os
import tempfile
import uuid

import matplotlib
matplotlib.use("Agg")  # GUI 없이 렌더링
import matplotlib.pyplot as plt
import koreanize_matplotlib
import pandas as pd
from langchain_core.tools import tool



# 차트 이미지 임시 저장 디렉토리
CHART_DIR = os.path.join(tempfile.gettempdir(), "streamlit_charts")
os.makedirs(CHART_DIR, exist_ok=True)

# 차트 경로 마커 (data_agent.py에서 파싱용)
CHART_MARKER = "CHART_PATH:"

# 생성된 차트 정보를 임시 보관 (data_agent.py에서 소비)
_pending_charts: list[dict] = []


def get_pending_charts() -> list[dict]:
    """생성된 차트 정보를 반환하고 내부 목록을 비웁니다."""
    charts = list(_pending_charts)
    _pending_charts.clear()
    return charts

# CSV 파일 경로 (프로젝트 루트 기준)
DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
CSV_PATH = os.path.join(DATA_DIR, "sample_sales.csv")


def _load_df() -> pd.DataFrame:
    """CSV 파일을 DataFrame으로 로드합니다 (내부 헬퍼)."""
    df = pd.read_csv(CSV_PATH)
    df["주문일"] = pd.to_datetime(df["주문일"])
    return df


def _save_chart(fig, title: str = "차트", chart_type: str = "bar") -> str:
    """matplotlib Figure를 PNG로 저장하고 차트 정보를 _pending_charts에 등록합니다."""
    filename = f"chart_{uuid.uuid4().hex[:8]}.png"
    filepath = os.path.join(CHART_DIR, filename)
    fig.savefig(filepath, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    _pending_charts.append({
        "title": title,
        "path": filepath,
        "type": chart_type,
    })
    return f"{CHART_MARKER}{filepath}"


# --- 데이터 조회 도구 ---


@tool
def load_csv_data() -> str:
    """CSV 데이터의 기본 정보를 요약하여 반환합니다.

    행 수, 컬럼 목록, 처음 5행 미리보기, 기본 통계를 포함합니다.
    데이터를 처음 살펴볼 때 이 도구를 사용하세요.
    """
    df = _load_df()

    info_parts = [
        f"[데이터 요약]",
        f"- 총 {len(df)}행, {len(df.columns)}개 컬럼",
        f"- 컬럼: {', '.join(df.columns.tolist())}",
        f"",
        f"[처음 5행 미리보기]",
        df.head().to_string(index=False),
        f"",
        f"[수치형 컬럼 기본 통계]",
        df.describe().to_string(),
    ]
    return "\n".join(info_parts)


@tool
def query_data(query: str) -> str:
    """Pandas 쿼리 코드를 실행하여 데이터를 분석합니다.

    Args:
        query: 실행할 Pandas 코드 문자열.
               변수 'df'로 DataFrame에 접근할 수 있습니다.
               예시: "df.groupby('카테고리')['총액'].sum().sort_values(ascending=False)"
               예시: "df[df['지역'] == '서울'].head(10)"
               예시: "df['카테고리'].value_counts()"

    Returns:
        쿼리 실행 결과를 문자열로 반환합니다.
    """
    df = _load_df()

    try:
        # 안전한 실행 환경 구성
        allowed_globals = {"__builtins__": {}}
        allowed_locals = {"df": df, "pd": pd}
        result = eval(query, allowed_globals, allowed_locals)

        if isinstance(result, pd.DataFrame):
            return f"결과 ({len(result)}행):\n{result.to_string(index=False)}"
        elif isinstance(result, pd.Series):
            return f"결과:\n{result.to_string()}"
        else:
            return f"결과: {result}"
    except Exception as e:
        return f"쿼리 실행 오류: {e}\n입력된 쿼리: {query}"


@tool
def get_column_stats(column_name: str) -> str:
    """특정 컬럼의 상세 통계 정보를 반환합니다.

    Args:
        column_name: 통계를 조회할 컬럼명.
                     사용 가능한 컬럼: 주문일, 카테고리, 상품명, 수량, 단가, 총액, 지역

    Returns:
        해당 컬럼의 통계 정보 (수치형: 평균/중앙값/최소/최대, 범주형: 고유값/빈도)
    """
    df = _load_df()

    if column_name not in df.columns:
        return f"[오류] '{column_name}' 컬럼이 존재하지 않습니다. 사용 가능한 컬럼: {', '.join(df.columns.tolist())}"

    col = df[column_name]
    parts = [f"'{column_name}' 컬럼 통계"]

    if pd.api.types.is_numeric_dtype(col):
        parts.extend([
            f"- 데이터 수: {col.count()}",
            f"- 평균: {col.mean():,.0f}",
            f"- 중앙값: {col.median():,.0f}",
            f"- 최솟값: {col.min():,.0f}",
            f"- 최댓값: {col.max():,.0f}",
            f"- 합계: {col.sum():,.0f}",
            f"- 표준편차: {col.std():,.0f}",
        ])
    else:
        value_counts = col.value_counts()
        parts.extend([
            f"- 고유값 수: {col.nunique()}",
            f"- 최빈값: {col.mode().iloc[0]}",
            f"",
            f"[빈도 분포]",
            value_counts.to_string(),
        ])

    return "\n".join(parts)


# --- 시각화 도구 ---

COLORS = ["#4E79A7", "#F28E2B", "#E15759", "#76B7B2", "#59A14F",
           "#EDC948", "#B07AA1", "#FF9DA7", "#9C755F", "#BAB0AC"]


@tool
def create_bar_chart(query: str, title: str, xlabel: str = "", ylabel: str = "") -> str:
    """데이터를 막대 차트로 시각화합니다.

    Args:
        query: Pandas 코드 문자열. 결과가 Series(index=카테고리, values=수치)여야 합니다.
               예시: "df.groupby('카테고리')['총액'].sum().sort_values(ascending=False)"
               예시: "df['지역'].value_counts()"
        title: 차트 제목 (예: "카테고리별 매출 합계")
        xlabel: X축 라벨 (선택)
        ylabel: Y축 라벨 (선택)

    Returns:
        차트 이미지 경로가 포함된 결과 문자열
    """
    df = _load_df()
    try:
        allowed_globals = {"__builtins__": {}}
        allowed_locals = {"df": df, "pd": pd}
        result = eval(query, allowed_globals, allowed_locals)

        if isinstance(result, pd.DataFrame):
            result = result.iloc[:, 0]  # 첫 번째 컬럼을 Series로

        fig, ax = plt.subplots(figsize=(10, 6))
        bars = ax.bar(result.index.astype(str), result.values, color=COLORS[:len(result)])

        # 값 라벨 표시
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2., height,
                    f'{height:,.0f}', ha='center', va='bottom', fontsize=10)

        ax.set_title(title, fontsize=16, fontweight="bold", pad=15)
        if xlabel:
            ax.set_xlabel(xlabel, fontsize=12)
        if ylabel:
            ax.set_ylabel(ylabel, fontsize=12)
        ax.tick_params(axis='x', rotation=45)
        fig.tight_layout()

        chart_path = _save_chart(fig, title=title, chart_type="bar")
        summary = f"막대 차트 생성 완료: {title}\n"
        summary += f"- 항목 수: {len(result)}개\n"
        summary += f"\n{chart_path}"
        return summary
    except Exception as e:
        return f"차트 생성 오류: {e}\n입력된 쿼리: {query}"


@tool
def create_line_chart(query: str, title: str, xlabel: str = "", ylabel: str = "") -> str:
    """데이터를 선(꺾은선) 차트로 시각화합니다. 시간 추이 분석에 적합합니다.

    Args:
        query: Pandas 코드 문자열. 결과가 Series(index=날짜/순서, values=수치)여야 합니다.
               예시: "df.groupby(df['주문일'].dt.to_period('M'))['총액'].sum()"
               예시: "df.groupby(df['주문일'].dt.date)['수량'].sum()"
        title: 차트 제목 (예: "월별 매출 추이")
        xlabel: X축 라벨 (선택)
        ylabel: Y축 라벨 (선택)

    Returns:
        차트 이미지 경로가 포함된 결과 문자열
    """
    df = _load_df()
    try:
        allowed_globals = {"__builtins__": {}}
        allowed_locals = {"df": df, "pd": pd}
        result = eval(query, allowed_globals, allowed_locals)

        if isinstance(result, pd.DataFrame):
            result = result.iloc[:, 0]

        fig, ax = plt.subplots(figsize=(10, 6))
        ax.plot(result.index.astype(str), result.values,
                color=COLORS[0], linewidth=2.5, marker="o", markersize=6)

        # 데이터 포인트 라벨
        for i, (x, y) in enumerate(zip(result.index.astype(str), result.values)):
            ax.annotate(f'{y:,.0f}', (x, y), textcoords="offset points",
                        xytext=(0, 10), ha='center', fontsize=9)

        ax.set_title(title, fontsize=16, fontweight="bold", pad=15)
        if xlabel:
            ax.set_xlabel(xlabel, fontsize=12)
        if ylabel:
            ax.set_ylabel(ylabel, fontsize=12)
        ax.tick_params(axis='x', rotation=45)
        ax.grid(True, alpha=0.3)
        fig.tight_layout()

        chart_path = _save_chart(fig, title=title, chart_type="line")
        summary = f"선 차트 생성 완료: {title}\n"
        summary += f"- 데이터 포인트: {len(result)}개\n"
        summary += f"\n{chart_path}"
        return summary
    except Exception as e:
        return f"차트 생성 오류: {e}\n입력된 쿼리: {query}"


@tool
def create_pie_chart(query: str, title: str) -> str:
    """데이터를 원형(파이) 차트로 시각화합니다. 비율/구성 분석에 적합합니다.

    Args:
        query: Pandas 코드 문자열. 결과가 Series(index=카테고리, values=수치)여야 합니다.
               예시: "df.groupby('카테고리')['총액'].sum()"
               예시: "df['지역'].value_counts()"
        title: 차트 제목 (예: "카테고리별 매출 비율")

    Returns:
        차트 이미지 경로가 포함된 결과 문자열
    """
    df = _load_df()
    try:
        allowed_globals = {"__builtins__": {}}
        allowed_locals = {"df": df, "pd": pd}
        result = eval(query, allowed_globals, allowed_locals)

        if isinstance(result, pd.DataFrame):
            result = result.iloc[:, 0]

        fig, ax = plt.subplots(figsize=(8, 8))
        wedges, texts, autotexts = ax.pie(
            result.values,
            labels=result.index.astype(str),
            autopct='%1.1f%%',
            colors=COLORS[:len(result)],
            startangle=90,
            textprops={'fontsize': 11},
        )
        for autotext in autotexts:
            autotext.set_fontsize(10)
            autotext.set_fontweight("bold")

        ax.set_title(title, fontsize=16, fontweight="bold", pad=20)
        fig.tight_layout()

        chart_path = _save_chart(fig, title=title, chart_type="pie")
        summary = f"파이 차트 생성 완료: {title}\n"
        summary += f"- 항목 수: {len(result)}개\n"
        summary += f"\n{chart_path}"
        return summary
    except Exception as e:
        return f"차트 생성 오류: {e}\n입력된 쿼리: {query}"


# --- 안전모 전용 차트 도구 ---

@tool
def create_helmet_compliance_chart(days: int = 14) -> str:
    """안전모 준수율 추이를 선(꺾은선) 차트로 시각화합니다.

    Args:
        days: 최근 며칠간의 데이터를 표시할지 (기본 14일)

    Returns:
        생성된 차트 이미지 경로가 포함된 결과 문자열
    """
    import helmet_data_manager
    df = helmet_data_manager.get_recent_days(days)

    if df.empty:
        return "데이터가 없어서 차트를 만들 수 없습니다."

    fig, ax = plt.subplots(figsize=(10, 5))
    dates = df["날짜"].astype(str)
    rates = df["준수율"].values

    ax.plot(dates, rates, color=COLORS[0], linewidth=2.5, marker="o", markersize=6)
    ax.axhline(y=90, color="red", linestyle="--", alpha=0.5, label="목표 준수율 90%")

    for i, (x, y) in enumerate(zip(dates, rates)):
        ax.annotate(f'{y}%', (x, y), textcoords="offset points",
                    xytext=(0, 10), ha='center', fontsize=9)

    ax.set_title(f"최근 {days}일 안전모 준수율 추이", fontsize=16, fontweight="bold", pad=15)
    ax.set_xlabel("날짜", fontsize=12)
    ax.set_ylabel("준수율 (%)", fontsize=12)
    ax.set_ylim(70, 105)
    ax.tick_params(axis='x', rotation=45)
    ax.grid(True, alpha=0.3)
    ax.legend()
    fig.tight_layout()

    chart_path = _save_chart(fig, title="준수율 추이", chart_type="line")
    return f"준수율 추이 차트 생성 완료 (최근 {days}일)\n{chart_path}"


@tool
def create_helmet_bar_chart(days: int = 14) -> str:
    """안전모 착용/미착용 인원을 막대 차트로 시각화합니다.

    Args:
        days: 최근 며칠간의 데이터를 표시할지 (기본 14일)

    Returns:
        생성된 차트 이미지 경로가 포함된 결과 문자열
    """
    import helmet_data_manager
    import numpy as np

    df = helmet_data_manager.get_recent_days(days)

    if df.empty:
        return "데이터가 없어서 차트를 만들 수 없습니다."

    fig, ax = plt.subplots(figsize=(10, 5))
    dates = df["날짜"].astype(str)
    x = np.arange(len(dates))
    width = 0.35

    bars1 = ax.bar(x - width/2, df["안전모착용"].values, width, label="착용", color=COLORS[0])
    bars2 = ax.bar(x + width/2, df["미착용"].values, width, label="미착용", color=COLORS[2])

    # 값 라벨
    for bar in bars1:
        ax.text(bar.get_x() + bar.get_width()/2., bar.get_height(),
                f'{int(bar.get_height())}', ha='center', va='bottom', fontsize=9)
    for bar in bars2:
        ax.text(bar.get_x() + bar.get_width()/2., bar.get_height(),
                f'{int(bar.get_height())}', ha='center', va='bottom', fontsize=9)

    ax.set_title(f"최근 {days}일 착용/미착용 인원", fontsize=16, fontweight="bold", pad=15)
    ax.set_xlabel("날짜", fontsize=12)
    ax.set_ylabel("인원 (명)", fontsize=12)
    ax.set_xticks(x)
    ax.set_xticklabels(dates, rotation=45)
    ax.legend()
    fig.tight_layout()

    chart_path = _save_chart(fig, title="착용/미착용 인원", chart_type="bar")
    return f"착용/미착용 막대 차트 생성 완료 (최근 {days}일)\n{chart_path}"


@tool
def generate_word_report(report_text: str, title: str = "주간 안전 진단 리포트") -> str:
    """분석 보고서를 Word(.docx) 파일로 생성합니다.

    보고서 텍스트와 함께, 이전에 생성된 차트 이미지들을 자동으로 삽입합니다.
    반드시 차트를 먼저 생성(create_helmet_compliance_chart, create_helmet_bar_chart)한 뒤
    이 도구를 호출하세요.

    Args:
        report_text: 보고서 본문 텍스트 (마크다운 형식 가능)
        title: 보고서 제목 (기본: "주간 안전 진단 리포트")

    Returns:
        생성된 Word 파일 경로
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import datetime

    doc = Document()

    # --- 제목 ---
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 날짜
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"작성일: {datetime.date.today().strftime('%Y-%m-%d')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")  # 빈 줄

    # --- 본문 ---
    # 마크다운 텍스트를 간단히 파싱해서 Word에 넣기
    import re
    img_pattern = re.compile(r'!\[([^\]]*)\]\((?:CHART_PATH:)?([^)]+)\)')

    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # 마크다운 이미지 구문 처리: ![alt](CHART_PATH:path) 또는 ![alt](path)
        img_match = img_pattern.match(line)
        if img_match:
            img_title = img_match.group(1)
            img_path = img_match.group(2)
            if os.path.exists(img_path):
                if img_title:
                    doc.add_heading(img_title, level=2)
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph("")  # 이미지 아래 여백
            continue

        # 이모지 제거하지 않고 그대로 넣기
        if line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("1.") or line.startswith("2.") or line.startswith("3.") or line.startswith("4."):
            doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph(line)

    # --- 차트 이미지 삽입 ---
    charts = get_pending_charts()
    if charts:
        doc.add_page_break()
        doc.add_heading("분석 차트", level=1)

        for chart_info in charts:
            chart_path = chart_info["path"]
            chart_title = chart_info["title"]

            if os.path.exists(chart_path):
                doc.add_heading(chart_title, level=2)
                doc.add_picture(chart_path, width=Inches(6))
                doc.add_paragraph("")  # 차트 아래 여백

    # --- 파일 저장 (프로젝트 안의 reports 폴더에 저장) ---
    REPORT_DIR = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(REPORT_DIR, exist_ok=True)
    filename = f"safety_report_{datetime.date.today().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6]}.docx"
    filepath = os.path.join(REPORT_DIR, filename)
    doc.save(filepath)

    return f"WORD_PATH:{filepath}"


# 에이전트에 등록할 도구 목록
ALL_TOOLS = [
    load_csv_data, query_data, get_column_stats,
    create_bar_chart, create_line_chart, create_pie_chart,
]

# 안전모 전용 도구 목록
HELMET_TOOLS = [
    create_helmet_compliance_chart,
    create_helmet_bar_chart,
    generate_word_report,
]


# ─────────────────────────────────────────────
# 급식 전용 도구 (Meal Analysis Tools)
# ─────────────────────────────────────────────

@tool
def get_meal_data(days: int = 7) -> str:
    """최근 N일간의 급식 현황 요약 텍스트를 반환합니다.

    Args:
        days: 조회할 일수 (기본 7일, 최대 31일)

    Returns:
        날짜별 칼로리·영양소 현황 텍스트
    """
    import meal_data_manager
    return meal_data_manager.get_summary_text(days)


@tool
def get_meal_statistics(days: int = 0) -> str:
    """급식 영양소 통계(평균·최소·최대·표준편차)를 반환합니다.

    Args:
        days: 최근 N일 필터 (0이면 전체 데이터)

    Returns:
        영양소별 상세 통계 텍스트
    """
    import meal_data_manager
    df = meal_data_manager.get_recent_days(days) if days > 0 else meal_data_manager.get_all_df()
    if df.empty:
        return "급식 데이터가 없습니다."

    nutrients = ["total_kcal", "total_carbs", "total_protein", "total_fat"]
    labels = {"total_kcal": "칼로리(kcal)", "total_carbs": "탄수화물(g)",
              "total_protein": "단백질(g)", "total_fat": "지방(g)"}
    lines = [f"[영양소 통계 - {'전체' if days == 0 else f'최근 {days}일'} {len(df)}일 기준]"]
    for col in nutrients:
        s = df[col]
        lines.append(
            f"  {labels[col]}: 평균 {s.mean():.1f} | 최솟값 {s.min():.1f}"
            f" | 최댓값 {s.max():.1f} | 표준편차 {s.std():.1f}"
        )
    return "\n".join(lines)


@tool
def get_food_intake_ranking() -> str:
    """음식별 평균 섭취율 순위를 반환합니다.

    섭취율이 높은 순서로 음식을 나열하여, 학생들이 선호하는 메뉴와
    잘 먹지 않는 메뉴를 파악할 수 있습니다.

    Returns:
        음식명, 평균 섭취율(%), 등장 횟수를 담은 텍스트
    """
    import meal_data_manager
    df = meal_data_manager.get_food_intake_df()
    if df.empty:
        return "음식 데이터가 없습니다."

    lines = ["[음식별 평균 섭취율 순위]"]
    for i, row in df.iterrows():
        lines.append(
            f"  {i+1}위. {row['food_name']:10s} | 평균 섭취율 {row['avg_intake_rate']:5.1f}% | 제공 횟수 {row['count']}회"
        )
    return "\n".join(lines)


@tool
def create_meal_kcal_chart(days: int = 14) -> str:
    """날짜별 급식 칼로리 추이 선(꺾은선) 차트를 생성합니다.

    Args:
        days: 최근 며칠간의 데이터를 표시할지 (기본 14일)

    Returns:
        차트 이미지 경로가 포함된 결과 문자열
    """
    import meal_data_manager
    df = meal_data_manager.get_recent_days(days)
    if df.empty:
        return "데이터가 없어서 차트를 만들 수 없습니다."

    fig, ax = plt.subplots(figsize=(11, 5))
    dates = df["날짜"].dt.strftime("%m/%d")
    kcal = df["total_kcal"].values

    ax.plot(dates, kcal, color=COLORS[0], linewidth=2.5, marker="o", markersize=7)
    ax.axhline(y=650, color="#E15759", linestyle="--", alpha=0.6, label="권장 칼로리 650kcal")

    for x, y in zip(dates, kcal):
        ax.annotate(f"{y}", (x, y), textcoords="offset points",
                    xytext=(0, 10), ha="center", fontsize=9)

    ax.set_title(f"최근 {days}일 급식 칼로리 추이", fontsize=16, fontweight="bold", pad=15)
    ax.set_xlabel("날짜", fontsize=12)
    ax.set_ylabel("칼로리 (kcal)", fontsize=12)
    ax.tick_params(axis="x", rotation=45)
    ax.grid(True, alpha=0.3)
    ax.legend()
    fig.tight_layout()

    chart_path = _save_chart(fig, title="날짜별 칼로리 추이", chart_type="line")
    return f"칼로리 추이 차트 생성 완료 (최근 {days}일)\n{chart_path}"


@tool
def create_meal_nutrition_bar_chart(days: int = 14) -> str:
    """날짜별 영양소(탄수화물·단백질·지방) 구성 막대 차트를 생성합니다.

    Args:
        days: 최근 며칠간의 데이터를 표시할지 (기본 14일)

    Returns:
        차트 이미지 경로가 포함된 결과 문자열
    """
    import numpy as np
    import meal_data_manager
    df = meal_data_manager.get_recent_days(days)
    if df.empty:
        return "데이터가 없어서 차트를 만들 수 없습니다."

    dates = df["날짜"].dt.strftime("%m/%d").tolist()
    x = np.arange(len(dates))
    width = 0.25

    fig, ax = plt.subplots(figsize=(12, 6))
    b1 = ax.bar(x - width, df["total_carbs"].values, width, label="탄수화물(g)", color=COLORS[0])
    b2 = ax.bar(x, df["total_protein"].values, width, label="단백질(g)", color=COLORS[3])
    b3 = ax.bar(x + width, df["total_fat"].values, width, label="지방(g)", color=COLORS[2])

    for bars in (b1, b2, b3):
        for bar in bars:
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2., h,
                    f"{h:.0f}", ha="center", va="bottom", fontsize=8)

    ax.set_title(f"최근 {days}일 영양소 구성", fontsize=16, fontweight="bold", pad=15)
    ax.set_xlabel("날짜", fontsize=12)
    ax.set_ylabel("g", fontsize=12)
    ax.set_xticks(x)
    ax.set_xticklabels(dates, rotation=45)
    ax.legend()
    fig.tight_layout()

    chart_path = _save_chart(fig, title="날짜별 영양소 구성", chart_type="bar")
    return f"영양소 구성 차트 생성 완료 (최근 {days}일)\n{chart_path}"


@tool
def create_food_intake_chart() -> str:
    """음식별 평균 섭취율 막대 차트를 생성합니다.

    섭취율이 낮은 음식을 한눈에 파악할 수 있어 식단 개선에 활용됩니다.

    Returns:
        차트 이미지 경로가 포함된 결과 문자열
    """
    import meal_data_manager
    df = meal_data_manager.get_food_intake_df()
    if df.empty:
        return "음식 데이터가 없습니다."

    fig, ax = plt.subplots(figsize=(12, max(6, len(df) * 0.4)))
    colors = [COLORS[0] if r >= 70 else COLORS[2] for r in df["avg_intake_rate"]]
    bars = ax.barh(df["food_name"], df["avg_intake_rate"], color=colors)

    for bar in bars:
        w = bar.get_width()
        ax.text(w + 1, bar.get_y() + bar.get_height() / 2.,
                f"{w:.1f}%", va="center", fontsize=9)

    ax.axvline(x=70, color="gray", linestyle="--", alpha=0.5, label="기준선 70%")
    ax.set_title("음식별 평균 섭취율", fontsize=16, fontweight="bold", pad=15)
    ax.set_xlabel("평균 섭취율 (%)", fontsize=12)
    ax.set_xlim(0, 115)
    ax.legend()
    fig.tight_layout()

    chart_path = _save_chart(fig, title="음식별 평균 섭취율", chart_type="bar")
    return f"음식별 섭취율 차트 생성 완료\n{chart_path}"


@tool
def generate_meal_report(report_text: str, title: str = "급식 분석 보고서") -> str:
    """급식 분석 보고서를 Word(.docx) 파일로 생성합니다.

    보고서 텍스트와 함께, 이전에 생성된 차트 이미지들을 자동으로 삽입합니다.
    반드시 차트를 먼저 생성한 뒤 이 도구를 호출하세요.

    Args:
        report_text: 보고서 본문 텍스트 (마크다운 형식 가능)
        title: 보고서 제목 (기본: "급식 분석 보고서")

    Returns:
        생성된 Word 파일 경로 (WORD_PATH: 접두사 포함)
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import datetime
    import re

    doc = Document()

    # --- 제목 ---
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"작성일: {datetime.date.today().strftime('%Y-%m-%d')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")

    # --- 본문 ---
    img_pattern = re.compile(r'!\[([^\]]*)\]\((?:CHART_PATH:)?([^)]+)\)')

    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        img_match = img_pattern.match(line)
        if img_match:
            img_title = img_match.group(1)
            img_path = img_match.group(2)
            if os.path.exists(img_path):
                if img_title:
                    doc.add_heading(img_title, level=2)
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[:2].isdigit() and line[1] == ".":
            doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph(line)

    # --- 차트 이미지 삽입 ---
    charts = get_pending_charts()
    if charts:
        doc.add_page_break()
        doc.add_heading("분석 차트", level=1)
        for chart_info in charts:
            if os.path.exists(chart_info["path"]):
                doc.add_heading(chart_info["title"], level=2)
                doc.add_picture(chart_info["path"], width=Inches(6))
                doc.add_paragraph("")

    # --- 저장 ---
    REPORT_DIR = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(REPORT_DIR, exist_ok=True)
    filename = f"meal_report_{datetime.date.today().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6]}.docx"
    filepath = os.path.join(REPORT_DIR, filename)
    doc.save(filepath)

    return f"WORD_PATH:{filepath}"


# 급식 전용 도구 목록
MEAL_TOOLS = [
    get_meal_data,
    get_meal_statistics,
    get_food_intake_ranking,
    create_meal_kcal_chart,
    create_meal_nutrition_bar_chart,
    create_food_intake_chart,
    generate_meal_report,
]
